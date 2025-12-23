import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup


def set_font_color(run, color=RGBColor(0, 0, 0)):
    """设置文本颜色为黑色"""
    run.font.color.rgb = color


def add_html_to_docx(doc, html_content):
    """将HTML内容添加到Word文档中"""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 遍历HTML中的所有元素
    for element in soup.children:
        if element.name == 'h1':
            heading = doc.add_heading(level=1)
            run = heading.add_run(element.text)
            set_font_color(run)
        elif element.name == 'h2':
            heading = doc.add_heading(level=2)
            run = heading.add_run(element.text)
            set_font_color(run)
        elif element.name == 'h3':
            heading = doc.add_heading(level=3)
            run = heading.add_run(element.text)
            set_font_color(run)
        elif element.name == 'h4':
            heading = doc.add_heading(level=4)
            run = heading.add_run(element.text)
            set_font_color(run)
        elif element.name == 'h5':
            heading = doc.add_heading(level=5)
            run = heading.add_run(element.text)
            set_font_color(run)
        elif element.name == 'h6':
            heading = doc.add_heading(level=6)
            run = heading.add_run(element.text)
            set_font_color(run)
        elif element.name == 'p':
            para = doc.add_paragraph()
            run = para.add_run(element.text)
            set_font_color(run)
        elif element.name == 'ul':
            # 处理无序列表
            for li in element.find_all('li'):
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(li.text)
                set_font_color(run)
        elif element.name == 'ol':
            # 处理有序列表
            for li in element.find_all('li'):
                para = doc.add_paragraph(style='List Number')
                run = para.add_run(li.text)
                set_font_color(run)
        elif element.name == 'table':
            # 处理表格
            rows = element.find_all('tr')
            if rows:
                # 确定表格列数
                first_row = rows[0]
                cols = max(len(first_row.find_all('th')), len(first_row.find_all('td')))
                
                # 创建表格
                table = doc.add_table(rows=0, cols=cols)
                table.style = 'Table Grid'
                
                # 填充表格内容
                for row in rows:
                    doc_row = table.add_row().cells
                    th_cells = row.find_all('th')
                    td_cells = row.find_all('td')
                    
                    # 优先使用表头单元格
                    cells = th_cells if th_cells else td_cells
                    
                    for i, cell in enumerate(cells):
                        if i < len(doc_row):
                            run = doc_row[i].paragraphs[0].add_run(cell.text)
                            set_font_color(run)
        elif element.name == 'pre':
            # 处理代码块
            para = doc.add_paragraph()
            run = para.add_run(element.text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            set_font_color(run)
        elif element.name == 'code':
            # 处理行内代码
            para = doc.add_paragraph()
            run = para.add_run(element.text)
            run.font.name = 'Courier New'
            set_font_color(run)
        elif element.name == 'br':
            # 处理换行
            doc.add_paragraph()


def md_to_docx(md_path, docx_path):
    # 创建Document对象
    doc = Document()
    
    # 读取Markdown文件内容
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 将Markdown转换为HTML
    html_content = markdown.markdown(md_content)
    
    # 将HTML内容添加到Word文档中
    add_html_to_docx(doc, html_content)
    
    # 保存Word文档
    doc.save(docx_path)
    print(f'Word文档已生成: {docx_path}')


if __name__ == '__main__':
    md_file = '系统完整报告.md'
    docx_file = '超市财务管理系统完整报告_new.docx'
    md_to_docx(md_file, docx_file)